# -*- coding: utf-8 -*-
"""
Gerador de Laudo Pericial v2.5 (Streamlit - Logo, Date/Time Consistent)

Este script gera laudos periciais para identificação de drogas e substâncias correlatas
usando o Streamlit com a logo da Polícia Científica e exibição de data/hora
consistente com o fuso horário de São Paulo. A opção de adicionar imagem
foi movida para após a descrição dos itens.

Requerimentos:
    - streamlit
    - python-docx
    - Pillow (PIL)
    - pytz

Uso:
    1. Instale as dependências: pip install streamlit python-docx Pillow pytz
    2. Salve este código como 'gerador_laudo.py'
    3. (Opcional) Salve a imagem do logo como 'logo_policia_cientifica.png' no mesmo diretório.
    4. Execute o script: streamlit run gerador_laudo.py
    5. Interaja com a interface web para gerar o laudo.
    6. Baixe o laudo gerado como um arquivo .docx.
"""

import re
from datetime import datetime
import io
import pytz # Importado para usar no relógio também
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches, RGBColor # Importar RGBColor explicitamente
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE # Importar WD_STYLE_TYPE
from PIL import Image
import time  # For the clock (embora não usado ativamente no loop principal)
# Importações necessárias para campos de página
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import traceback # Para mostrar erros detalhados

# --- Constantes ---
TIPOS_MATERIAL_BASE = {
    "v": "vegetal dessecado",
    "po": "pulverizado",
    "pd": "petrificado",
    "r": "resinoso"
}

TIPOS_EMBALAGEM_BASE = {
    "e": "microtubo do tipo eppendorf",
    "z": "embalagem do tipo ziplock",
    "a": "papel alumínio",
    "pl": "plástico",
    "pa": "papel"
}

CORES_FEMININO_EMBALAGEM = {
    "t": "transparente", "b": "branca", "az": "azul", "am": "amarela",
    "vd": "verde", "vm": "vermelha", "p": "preta", "c": "cinza",
    "m": "marrom", "r": "rosa", "l": "laranja", "violeta": "violeta", "roxa": "roxa"
}

QUANTIDADES_EXTENSO = {
    1: "uma", 2: "duas", 3: "três", 4: "quatro", 5: "cinco",
    6: "seis", 7: "sete", 8: "oito", 9: "nove", 10: "dez"
}

# Dicionário de meses para garantir português
meses_portugues = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril", 5: "maio", 6: "junho",
    7: "julho", 8: "agosto", 9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
}

# Mapeamento dias da semana (datetime.weekday() -> 0=Segunda)
dias_semana_portugues = {
    0: "Segunda-feira", 1: "Terça-feira", 2: "Quarta-feira", 3: "Quinta-feira",
    4: "Sexta-feira", 5: "Sábado", 6: "Domingo"
}


# --- Funções Auxiliares ---
def pluralizar_palavra(palavra, quantidade):
    """Pluraliza palavras em português (com algumas regras básicas)."""
    if quantidade == 1:
        return palavra
    # Casos especiais que não pluralizam ou têm forma específica
    if palavra in ["microtubo do tipo eppendorf", "embalagem do tipo ziplock", "papel alumínio"]:
        return palavra
    if palavra.endswith('m') and palavra not in ["alumínio"]: # Exceção para alumínio
        return re.sub(r'm$', 'ns', palavra) # marrom -> marrons
    if palavra.endswith('ão'):
        return re.sub(r'ão$', 'ões', palavra) # porção -> porções
    elif palavra.endswith(('r', 'z', 's')): # Acrescenta 'es' para r, z, s (pulverizado->pulverizados?) - Revisar regra 's'
         # Se termina em 's' e é paroxítona ou proparoxítona, não muda (ex: lápis). Se oxítona, +es (ex: gás->gases)
         # Simplificação: Adicionar 'es' para r, z. Manter 's' para palavras terminadas em 's'.
         if palavra.endswith(('r', 'z')):
              return palavra + 'es' # cor -> cores
         else: # Termina em 's'
              return palavra # plástico -> plásticos (já termina em s)
    elif palavra.endswith('l'):
         # Troca 'l' por 'is' (papel -> papéis, vegetal -> vegetais, azul -> azuis)
         return palavra[:-1] + 'is'
    # Regra geral: adiciona 's'
    else:
        # Ex: petrificado -> petrificados, resinoso -> resinosos, branca -> brancas
        return palavra + 's'

def obter_quantidade_extenso(qtd):
    """Retorna a quantidade por extenso (1-10) ou o número como string."""
    return QUANTIDADES_EXTENSO.get(qtd, str(qtd))

def adicionar_paragrafo(doc, text, style=None, align=None, color=None, size=None, bold=False, italic=False):
    """Adiciona um parágrafo ao documento docx com formatação flexível."""
    p = doc.add_paragraph() # Sempre adiciona um parágrafo base

    # Aplica estilo de parágrafo, se fornecido e existir
    if style and style in doc.styles:
         p.style = doc.styles[style]
    elif style:
         # Aviso opcional se o estilo não for encontrado
         # print(f"Aviso: Estilo '{style}' não encontrado no documento.")
         p.style = doc.styles['Normal'] # Aplica estilo Normal como fallback

    # Aplica alinhamento
    if align:
        align_map = {
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'left': WD_ALIGN_PARAGRAPH.LEFT # Default já é LEFT
        }
        # Usa get com fallback para LEFT se o alinhamento for inválido
        p.alignment = align_map.get(str(align).lower(), WD_ALIGN_PARAGRAPH.LEFT)

    # Adiciona o texto com formatação de caractere
    run = p.add_run(text)
    if color:
        # Assume que color é um RGBColor object ou uma tupla/lista (R, G, B)
        try:
            if isinstance(color, RGBColor):
                 run.font.color.rgb = color
            elif isinstance(color, (tuple, list)) and len(color) == 3:
                 run.font.color.rgb = RGBColor(color[0], color[1], color[2])
            # else: print(f"Aviso: Formato de cor inválido: {color}") # Opcional
        except Exception as e:
            print(f"Erro ao aplicar cor: {e}")
    if size:
        try:
            run.font.size = Pt(int(size))
        except ValueError:
            print(f"Tamanho de fonte inválido: {size}")
    if bold:
        run.font.bold = True # True ou False
    if italic:
        run.font.italic = True # True ou False

def inserir_imagem_docx(doc, image_file_uploader):
    """Insere uma imagem vinda do st.file_uploader no documento docx, centralizada."""
    try:
        if image_file_uploader:
            # Ler a imagem para um stream de bytes para evitar problemas com tipos de arquivo
            img_stream = io.BytesIO(image_file_uploader.getvalue())
            img = Image.open(img_stream)
            width_px, height_px = img.size

            # Define a largura máxima em polegadas (e.g., 6 polegadas para A4 com margens de 1")
            max_width_inches = 6.0
            # Converte DPI para pixels por polegada (assumindo 96 DPI como padrão comum)
            # É melhor usar a resolução da imagem se disponível, mas 96 é um fallback razoável
            dpi = img.info.get('dpi', (96, 96))[0]
            if dpi == 0: dpi = 96 # Evita divisão por zero

            width_inches = width_px / dpi
            # height_inches = height_px / dpi # Calculado mas não usado diretamente abaixo

            # Redimensiona se exceder a largura máxima, mantendo a proporção
            if width_inches > max_width_inches:
                # Calcula a nova altura proporcionalmente
                aspect_ratio = height_px / width_px
                display_width_inches = max_width_inches
                display_height_inches = display_width_inches * aspect_ratio
            else:
                display_width_inches = width_inches
                # A altura será ajustada automaticamente pelo Word se passarmos só a largura
                # display_height_inches = height_inches

            # Adiciona a imagem centralizada
            # Adiciona um parágrafo específico para a imagem com alinhamento central
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            # Passa o stream de bytes original
            img_stream.seek(0) # Reseta o ponteiro do stream
            # Adiciona a imagem especificando apenas a largura; a altura se ajusta
            run.add_picture(img_stream, width=Inches(display_width_inches))

    except Exception as e:
        st.error(f"Erro ao inserir imagem no docx: {e}")
        # Log para console para debug
        print(f"Erro detalhado ao inserir imagem: {e}\n{traceback.format_exc()}")


def configurar_estilos(doc):
    """Configura os estilos de parágrafo e caractere do documento docx."""
    # --- PALETA DE CORES PADRÃO (TEXTO PRETO EM FUNDO BRANCO) ---
    # Para maior compatibilidade com diferentes versões do Word e leitores.
    COR_TEXTO_PRINCIPAL = RGBColor(0x00, 0x00, 0x00) # Preto padrão
    COR_DESTAQUE = RGBColor(0x00, 0x00, 0x00) # Preto para títulos secundários também
    COR_TEXTO_SECUNDARIO = RGBColor(0x59, 0x59, 0x59) # Cinza escuro

    # --- SE QUISER TENTAR A PALETA ESCURA (Pode não funcionar bem com fundo) ---
    # COR_TEXTO_PRINCIPAL = RGBColor(0xE0, 0xE0, 0xE0) # Cinza muito claro (quase branco)
    # COR_DESTAQUE = RGBColor(0x58, 0xA6, 0xFF)       # Azul claro (semelhante ao link do GitHub)
    # COR_TEXTO_SECUNDARIO = RGBColor(0x8B, 0x94, 0x9E) # Cinza médio
    # NOTA: O fundo da página NÃO será alterado por esta função.

    # Verifica se os estilos já existem para evitar erros, senão adiciona
    def get_or_add_style(doc, style_name, style_type):
         if style_name in doc.styles:
             return doc.styles[style_name]
         else:
             try:
                 return doc.styles.add_style(style_name, style_type)
             except Exception as e:
                  print(f"Erro ao adicionar estilo '{style_name}': {e}. Usando 'Normal'.")
                  return doc.styles['Normal'] # Fallback para Normal

    # Estilo 'Normal' (base para muitos outros)
    paragrafo_style = doc.styles['Normal']
    paragrafo_style.font.name = 'Calibri' # Fonte padrão e comum
    paragrafo_style.font.size = Pt(12)
    paragrafo_style.font.color.rgb = COR_TEXTO_PRINCIPAL
    paragrafo_style.paragraph_format.line_spacing = 1.15 # Espaçamento entre linhas (1.0 para simples)
    paragrafo_style.paragraph_format.space_before = Pt(0) # Sem espaço antes por padrão
    paragrafo_style.paragraph_format.space_after = Pt(8) # Espaço padrão após parágrafo (ajustar conforme gosto)

    # Estilo para o título principal (ex: HISTÓRICO)
    titulo_principal_style = get_or_add_style(doc, 'TituloPrincipal', WD_STYLE_TYPE.PARAGRAPH)
    titulo_principal_style.base_style = doc.styles['Normal'] # Baseado no Normal
    titulo_principal_style.font.name = 'Calibri'
    titulo_principal_style.font.size = Pt(14)
    titulo_principal_style.font.bold = True
    titulo_principal_style.font.color.rgb = COR_DESTAQUE # Usa a cor de destaque
    titulo_principal_style.paragraph_format.space_before = Pt(12) # Mais espaço antes
    titulo_principal_style.paragraph_format.space_after = Pt(6)  # Menos espaço depois

    # Estilo para títulos secundários (ex: Material Recebido)
    titulo_secundario_style = get_or_add_style(doc, 'TituloSecundario', WD_STYLE_TYPE.PARAGRAPH)
    titulo_secundario_style.base_style = doc.styles['Normal']
    titulo_secundario_style.font.name = 'Calibri'
    titulo_secundario_style.font.size = Pt(12)
    titulo_secundario_style.font.bold = True
    titulo_secundario_style.font.color.rgb = COR_DESTAQUE # Mesma cor do principal ou outra
    titulo_secundario_style.paragraph_format.space_before = Pt(10)
    titulo_secundario_style.paragraph_format.space_after = Pt(4)

    # Estilo para texto itálico (estilo de caractere)
    # Tenta obter, se não existir, cria baseado no Default Character Font
    if 'Italico' not in doc.styles:
         try:
            italico_style = doc.styles.add_style('Italico', WD_STYLE_TYPE.CHARACTER)
            italico_style.font.italic = True
            # Herda outras propriedades do Default Character Font
            italico_style.base_style = doc.styles['Default Paragraph Font'] # Ou pode deixar sem base
         except:
             print("Não foi possível criar estilo 'Italico'.")
    else:
         italico_style = doc.styles['Italico']
         italico_style.font.italic = True # Garante que está em itálico


    # Estilo para legendas de ilustrações
    ilustracao_style = get_or_add_style(doc, 'Ilustracao', WD_STYLE_TYPE.PARAGRAPH)
    ilustracao_style.base_style = doc.styles['Normal']
    ilustracao_style.font.name = 'Calibri'
    ilustracao_style.font.size = Pt(10) # Menor
    ilustracao_style.font.bold = False # Sem negrito
    ilustracao_style.font.italic = True # Itálico para legenda
    ilustracao_style.font.color.rgb = COR_TEXTO_SECUNDARIO # Cor secundária (cinza)
    ilustracao_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # Centralizado
    ilustracao_style.paragraph_format.space_before = Pt(4) # Pouco espaço antes
    ilustracao_style.paragraph_format.space_after = Pt(10) # Mais espaço depois


def configurar_pagina(doc):
    """Configura margens da página (padrão ABNT)."""
    for section in doc.sections:
        section.page_height = Inches(11.69) # A4 Altura
        section.page_width = Inches(8.27)  # A4 Largura
        section.top_margin = Inches(1.18)    # 3 cm
        section.bottom_margin = Inches(0.79) # 2 cm
        section.left_margin = Inches(1.18)   # 3 cm
        section.right_margin = Inches(0.79)  # 2 cm

def adicionar_cabecalho_rodape(doc):
    """Adiciona cabeçalho e rodapé padrão ao documento docx."""
    section = doc.sections[0] # Aplica à primeira seção

    # Cabeçalho (Ex: Nome da Instituição ou Tipo de Documento)
    header = section.header
    # Limpa cabeçalho existente, se houver (parágrafo padrão vazio)
    if header.paragraphs:
        # Remove todos os parágrafos existentes no cabeçalho
        for para in header.paragraphs:
             p_element = para._element
             p_element.getparent().remove(p_element)
             # header._body.remove(p_element) # Alternativa

    # Adiciona novo parágrafo no cabeçalho
    header_paragraph = header.add_paragraph()
    # Adiciona texto. Usar tab (\t) pode ser instável. Alinhamento é melhor.
    run_header = header_paragraph.add_run("POLÍCIA CIENTÍFICA DE GOIÁS") # Exemplo
    run_header.font.name = 'Calibri'
    run_header.font.size = Pt(10)
    run_header.font.bold = True
    # Adiciona um segundo run com o tipo de laudo à direita
    header_paragraph.add_run("\t\t\tLAUDO DE PERÍCIA CRIMINAL").bold = False # Tenta alinhar com tabs
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT # Alinha todo o parágrafo à esquerda (ou RIGHT/CENTER)
    # Para ter elementos à esquerda e direita, usar tabelas no cabeçalho é mais robusto

    # Rodapé (Ex: Endereço e Número de Página)
    footer = section.footer
    # Limpa rodapé existente
    if footer.paragraphs:
        for para in footer.paragraphs:
             p_element = para._element
             p_element.getparent().remove(p_element)

    footer_paragraph = footer.add_paragraph()
    # Exemplo de endereço (alinhado à esquerda)
    # run_footer_addr = footer_paragraph.add_run("Endereço do Instituto de Criminalística...")
    # run_footer_addr.font.size = Pt(9)

    # Número da página (alinhado à direita ou centro)
    # Usar um parágrafo separado para o número da página simplifica o alinhamento
    page_num_paragraph = footer.add_paragraph()
    page_num_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER # Ou RIGHT

    run_page = page_num_paragraph.add_run("Página ")
    run_page.font.name = 'Calibri'
    run_page.font.size = Pt(10)

    # Adiciona o campo de número da página (PAGE)
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    run_page._r.append(fld_char_begin)

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE \* MERGEFORMAT' # Código do campo PAGE
    run_page._r.append(instr_text)

    fld_char_sep = OxmlElement('w:fldChar')
    fld_char_sep.set(qn('w:fldCharType'), 'separate')
    run_page._r.append(fld_char_sep)

    # Adiciona um valor inicial (opcional, Word atualiza)
    # run_t = OxmlElement('w:t')
    # run_t.text = '1'
    # run_page._r.append(run_t)

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run_page._r.append(fld_char_end)

    # Adiciona " de " e o número total de páginas (NUMPAGES)
    run_num_pages = page_num_paragraph.add_run(" de ")
    run_num_pages.font.name = 'Calibri'
    run_num_pages.font.size = Pt(10)

    fld_char_begin_np = OxmlElement('w:fldChar')
    fld_char_begin_np.set(qn('w:fldCharType'), 'begin')
    run_num_pages._r.append(fld_char_begin_np)

    instr_text_np = OxmlElement('w:instrText')
    instr_text_np.set(qn('xml:space'), 'preserve')
    instr_text_np.text = 'NUMPAGES \* MERGEFORMAT' # Código do campo NUMPAGES
    run_num_pages._r.append(instr_text_np)

    fld_char_sep_np = OxmlElement('w:fldChar')
    fld_char_sep_np.set(qn('w:fldCharType'), 'separate')
    run_num_pages._r.append(fld_char_sep_np)

    fld_char_end_np = OxmlElement('w:fldChar')
    fld_char_end_np.set(qn('w:fldCharType'), 'end')
    run_num_pages._r.append(fld_char_end_np)


def adicionar_preambulo(doc, dados_laudo):
    """Adiciona o preâmbulo/histórico ao laudo."""
    adicionar_paragrafo(doc, "1 HISTÓRICO", style='TituloPrincipal', align='left')
    # Obter dados do dicionário, com valores padrão se não encontrados
    hoje = datetime.now(pytz.timezone('America/Sao_Paulo'))
    dia_atual = hoje.day
    mes_atual = meses_portugues.get(hoje.month, "Mês Inválido")
    ano_atual = hoje.year
    autoridade = dados_laudo.get('autoridade', '[Autoridade Requisitante Não Informada]')
    numero_doc = dados_laudo.get('num_doc', '[Documento Não Informado]')
    boletim = dados_laudo.get('boletim', '[BO Não Informado]')
    perito = "Daniel Chendes Lima" # Ou obter de dados_laudo se for variável

    texto_preambulo = (
        f"Aos {dia_atual} dias do mês de {mes_atual} de {ano_atual}, neste Instituto de Criminalística, "
        f"em atendimento à requisição formulada pelo(a) {autoridade}, constante do(a) {numero_doc}, "
        f"referente ao(à) {boletim}, compareceu a este setor o Perito Criminal {perito}, signatário, "
        "a fim de proceder aos exames periciais no material descrito a seguir."
    )
    adicionar_paragrafo(doc, texto_preambulo, align='justify', style='Normal')


def adicionar_material_recebido(doc, dados_laudo):
    """Adiciona a seção '2 MATERIAL RECEBIDO PARA EXAME' ao laudo docx."""
    adicionar_paragrafo(doc, "2 MATERIAL RECEBIDO PARA EXAME", style='TituloPrincipal')
    # Adicionar parágrafo sobre acondicionamento geral, se aplicável
    # Pode ser mais específico se a informação for coletada
    adicionar_paragrafo(doc,
        "O material foi recebido neste Instituto devidamente acondicionado e lacrado, "
        "preservando-se a cadeia de custódia.", # Texto genérico
        align='justify', style='Normal')

    # Adicionar a ilustração aqui, se houver
    imagem_carregada = dados_laudo.get('imagem') # Pega o objeto FileUploader do estado
    if imagem_carregada:
         inserir_imagem_docx(doc, imagem_carregada) # Passa o objeto para a função
         # Adiciona legenda após a imagem
         adicionar_paragrafo(doc, "Ilustração 1: Material(is) recebido(s) conforme encaminhado(s).", style='Ilustracao')

    subitens_cannabis = {}
    subitens_cocaina = {}

    if not dados_laudo.get('itens'): # Verifica se a lista de itens existe e não está vazia
         adicionar_paragrafo(doc, "Nenhum item de material foi descrito para exame.", style='Normal')
         return subitens_cannabis, subitens_cocaina

    for i, item in enumerate(dados_laudo['itens']):
        qtd = item.get('qtd', 1)
        qtd_ext = obter_quantidade_extenso(qtd)
        tipo_mat_cod = item.get('tipo_mat', '')
        tipo_material = TIPOS_MATERIAL_BASE.get(tipo_mat_cod, f"tipo '{tipo_mat_cod}'") # Mantém código se desconhecido

        emb_cod = item.get('emb', '')
        embalagem = TIPOS_EMBALAGEM_BASE.get(emb_cod, f"embalagem '{emb_cod}'")

        # Cor da embalagem
        cor_emb_cod = item.get('cor_emb') # Pode ser None ou ''
        desc_cor = ""
        if cor_emb_cod and emb_cod in ['pl', 'pa', 'e', 'z']: # Verifica se a cor é aplicável
            cor = CORES_FEMININO_EMBALAGEM.get(cor_emb_cod, cor_emb_cod) # Usa código se não mapeado
            # Ajuste para gênero da cor/embalagem (simplificado)
            # Assume que a maioria das cores funciona com "de cor X"
            desc_cor = f" de cor {cor}"

        # Pluralização
        # Pluraliza a embalagem primeiro, depois adiciona a cor se houver
        embalagem_base_plural = pluralizar_palavra(embalagem, qtd)
        embalagem_final = f"{embalagem_base_plural}{desc_cor}" # Junta

        porcao = pluralizar_palavra("porção", qtd)
        acond = "acondicionada em" if qtd == 1 else "acondicionadas individualmente em"

        ref_texto = f", relacionada a {item['pessoa']}" if item.get('pessoa') else ""
        subitem_ref = item.get('ref', '[Ref. Constatação Não Informada]')
        final_ponto = "."

        texto = (f"2.{i + 1} – {qtd} ({qtd_ext}) {porcao} de material {tipo_material}, "
                 f"{acond} {embalagem_final}, referente(s) à(s) amostra(s) "
                 f"do(s) subitem(ns) {subitem_ref} do laudo de constatação supracitado"
                 f"{ref_texto}{final_ponto}")
        adicionar_paragrafo(doc, texto, style='Normal', align='justify')

        # Mapeamento para seções de exames/resultados
        if tipo_mat_cod in ["v", "r"]:
            subitens_cannabis[subitem_ref if subitem_ref else f"Item_2.{i+1}"] = f"2.{i + 1}"
        elif tipo_mat_cod in ["po", "pd"]:
            subitens_cocaina[subitem_ref if subitem_ref else f"Item_2.{i+1}"] = f"2.{i + 1}"

    return subitens_cannabis, subitens_cocaina


def adicionar_objetivo_exames(doc):
    """Adiciona a seção '3 OBJETIVO DOS EXAMES'."""
    adicionar_paragrafo(doc, "3 OBJETIVO DOS EXAMES", style='TituloPrincipal')
    texto = ("O objetivo dos exames é identificar a natureza do material apresentado, "
             "verificando a presença de substâncias entorpecentes ou de uso proscrito no Brasil, "
             "listadas na Portaria SVS/MS nº 344/1998 e suas atualizações, a fim de "
             "constatar a materialidade de eventual infração penal.")
    adicionar_paragrafo(doc, texto, align='justify', style='Normal')

def adicionar_exames(doc, subitens_cannabis, subitens_cocaina):
    """Adiciona a seção '4 EXAMES'."""
    adicionar_paragrafo(doc, "4 EXAMES", style='TituloPrincipal')
    adicionar_paragrafo(doc,
     "Os materiais recebidos foram submetidos aos seguintes exames e testes:",
     style='Normal', align='justify')

    has_cannabis_item = bool(subitens_cannabis)
    has_cocaina_item = bool(subitens_cocaina)
    idx_counter = 1 # Contador para numeração dos tipos de exame

    # Adicionar Exame Macroscópico sempre que houver material
    if subitens_cannabis or subitens_cocaina or dados_laudo.get('itens'): # Se houver qqr item
         adicionar_paragrafo(doc, f"4.{idx_counter} Exame macroscópico:", style='TituloSecundario')
         adicionar_paragrafo(doc, "Observação das características gerais do material (aspecto, cor, odor, etc.).", style='Normal')
         idx_counter += 1


    if has_cannabis_item:
        adicionar_paragrafo(doc, f"4.{idx_counter} Testes para identificação de Cannabis sativa L.:", style='TituloSecundario')
        adicionar_paragrafo(doc, f"   a) Reação com Duquenois-Levine modificado;", style='Normal')
        adicionar_paragrafo(doc, f"   b) Reação com Fast Blue B Salt (Sal de Azul Sólido B);", style='Normal')
        adicionar_paragrafo(doc, f"   c) Cromatografia em Camada Delgada (CCD) comparativa com padrão analítico.", style='Normal')
        idx_counter += 1

    if has_cocaina_item:
        adicionar_paragrafo(doc, f"4.{idx_counter} Testes para identificação de cocaína:", style='TituloSecundario')
        adicionar_paragrafo(doc, f"   a) Reação com Tiocianato de Cobalto;", style='Normal')
        adicionar_paragrafo(doc, f"   b) Cromatografia em Camada Delgada (CCD) comparativa com padrão analítico.", style='Normal')
        idx_counter += 1

    # Se houver outros tipos de material não cobertos, adicionar menção genérica
    itens_outros = [item for item in dados_laudo.get('itens', []) if item.get('tipo_mat') not in ["v", "r", "po", "pd"]]
    if itens_outros and not has_cannabis_item and not has_cocaina_item: # Se SÓ tem outros
         adicionar_paragrafo(doc, f"4.{idx_counter} Testes para identificação de outras substâncias:", style='TituloSecundario')
         adicionar_paragrafo(doc, "Realização de testes químicos colorimétricos e/ou técnicas instrumentais apropriadas para a suspeita levantada (não detalhados neste modelo).", style='Normal')
         idx_counter += 1


def adicionar_resultados(doc, subitens_cannabis, subitens_cocaina, dados_laudo):
    """Adiciona a seção '5 RESULTADOS'."""
    adicionar_paragrafo(doc, "5 RESULTADOS", style='TituloPrincipal')
    idx_counter = 1 # Contador para os itens de resultado

    itens_cannabis = [item for item in dados_laudo.get('itens', []) if item.get('tipo_mat') in ["v", "r"]]
    itens_cocaina = [item for item in dados_laudo.get('itens', []) if item.get('tipo_mat') in ["po", "pd"]]

    if itens_cannabis:
        # Junta as referências dos itens (ex: "2.1, 2.3")
        desc_itens_str = ", ".join(sorted(subitens_cannabis.values()))
        label_desc = "no item" if len(subitens_cannabis) == 1 else "nos itens"

        adicionar_paragrafo(doc, f"5.{idx_counter} Para o(s) material(is) {label_desc} {desc_itens_str}:", style='TituloSecundario')
        adicionar_paragrafo(doc, "   a) Exame macroscópico: Material vegetal dessecado/resinoso com odor e características morfológicas compatíveis com Cannabis sativa L. (maconha).", style='Normal')
        adicionar_paragrafo(doc, "   b) Testes químicos (Duquenois-Levine e Fast Blue B Salt): Resultados positivos para canabinoides.", style='Normal')
        adicionar_paragrafo(doc, "   c) Cromatografia em Camada Delgada (CCD): Revelou manchas com Rf (Fator de Retenção) e coloração compatíveis com o padrão de Tetrahidrocanabinol (THC).", style='Normal')
        idx_counter += 1

    if itens_cocaina:
        desc_itens_str = ", ".join(sorted(subitens_cocaina.values()))
        label_desc = "no item" if len(subitens_cocaina) == 1 else "nos itens"

        adicionar_paragrafo(doc, f"5.{idx_counter} Para o(s) material(is) {label_desc} {desc_itens_str}:", style='TituloSecundario')
        adicionar_paragrafo(doc, "   a) Exame macroscópico: Material pulverulento/petrificado de coloração branca/amarelada, inodoro ou com odor característico.", style='Normal') # Texto exemplo
        adicionar_paragrafo(doc, "   b) Teste químico (Tiocianato de Cobalto): Resultado positivo para cocaína.", style='Normal')
        adicionar_paragrafo(doc, "   c) Cromatografia em Camada Delgada (CCD): Revelou mancha com Rf (Fator de Retenção) e coloração compatível com o padrão de Cloridrato de Cocaína.", style='Normal')
        idx_counter += 1

    # Resultados para outros materiais (genérico)
    itens_outros = [item for item in dados_laudo.get('itens', []) if item.get('tipo_mat') not in ["v", "r", "po", "pd"]]
    if itens_outros:
         desc_itens_str = ", ".join(sorted([f"2.{i+1}" for i, item in enumerate(dados_laudo['itens']) if item.get('tipo_mat') not in ["v", "r", "po", "pd"]]))
         label_desc = "no item" if len(itens_outros) == 1 else "nos itens"
         adicionar_paragrafo(doc, f"5.{idx_counter} Para o(s) material(is) {label_desc} {desc_itens_str}:", style='TituloSecundario')
         adicionar_paragrafo(doc, "   a) Exame macroscópico: [Descrever características observadas].", style='Normal')
         adicionar_paragrafo(doc, "   b) Demais testes: [Relatar resultados positivos ou negativos para outras substâncias, se aplicável, ou indicar 'negativo para as substâncias pesquisadas'].", style='Normal')
         idx_counter += 1

    if idx_counter == 1: # Nenhum resultado adicionado (sem itens válidos?)
        adicionar_paragrafo(doc, "Nenhum resultado a relatar para os itens fornecidos.", style='Normal')


def adicionar_conclusao(doc, dados_laudo):
    """Adiciona a seção '6 CONCLUSÃO'."""
    adicionar_paragrafo(doc, "6 CONCLUSÃO", style='TituloPrincipal')
    conclusoes = []
    itens_cannabis = [item for item in dados_laudo.get('itens', []) if item.get('tipo_mat') in ["v", "r"]]
    itens_cocaina = [item for item in dados_laudo.get('itens', []) if item.get('tipo_mat') in ["po", "pd"]]
    itens_outros = [item for item in dados_laudo.get('itens', []) if item.get('tipo_mat') not in ["v", "r", "po", "pd"]]


    # Referência legal padrão
    ref_legal = ("Portaria SVS/MS nº 344/1998 e suas atualizações")

    if itens_cannabis:
        desc_itens_nums = sorted([f"2.{i+1}" for i, item in enumerate(dados_laudo['itens']) if item.get('tipo_mat') in ["v", "r"]])
        desc_str = ", ".join(desc_itens_nums)
        label_desc = "no material descrito no item" if len(desc_itens_nums) == 1 else "nos materiais descritos nos itens"

        conclusoes.append(
            f"{label_desc} {desc_str}, foi detectada a presença de Tetrahidrocanabinol (THC), principal componente psicoativo da planta "
            f"Cannabis sativa L. (maconha). A planta Cannabis sativa L. e o THC são substâncias de uso proscrito no Brasil, "
            f"constantes na lista F1 (plantas) e F2 (substâncias psicotrópicas) da {ref_legal}, respectivamente."
        )

    if itens_cocaina:
        desc_itens_nums = sorted([f"2.{i+1}" for i, item in enumerate(dados_laudo['itens']) if item.get('tipo_mat') in ["po", "pd"]])
        desc_str = ", ".join(desc_itens_nums)
        label_desc = "no material descrito no item" if len(desc_itens_nums) == 1 else "nos materiais descritos nos itens"

        conclusoes.append(
            f"{label_desc} {desc_str}, foi detectada a presença de Cocaína. A Cocaína é uma substância "
            f"estimulante do sistema nervoso central, de uso proscrito no Brasil, constante na lista F2 (substâncias psicotrópicas) da {ref_legal}."
        )

    if itens_outros:
         desc_itens_nums = sorted([f"2.{i+1}" for i, item in enumerate(dados_laudo['itens']) if item.get('tipo_mat') not in ["v", "r", "po", "pd"]])
         desc_str = ", ".join(desc_itens_nums)
         label_desc = "no material descrito no item" if len(itens_outros) == 1 else "nos materiais descritos nos itens"
         # Conclusão para outros depende do resultado (positivo para outra droga ou negativo)
         conclusoes.append(
             f"{label_desc} {desc_str}, [concluir sobre a presença ou ausência de outras substâncias pesquisadas ou "
             f"relatar resultado negativo para maconha e cocaína, conforme o caso]."
             # Exemplo negativo: "não foram detectadas as substâncias maconha ou cocaína."
             # Exemplo positivo: "foi detectada a presença de [Nome da Substância], substância constante na lista [X] da Portaria..."
         )


    if conclusoes:
        texto_final = "Face ao exposto e com base nos resultados obtidos, conclui-se que "
        if len(conclusoes) > 1:
             # Junta as conclusões com ponto e vírgula, exceto a última que usa 'e'
             texto_final += "; ".join(conclusoes[:-1])
             texto_final += "; e " + conclusoes[-1]
        else:
             texto_final += conclusoes[0]
        texto_final = texto_final.replace("..", ".").replace(".,", ",") # Limpeza final
    elif dados_laudo.get('itens'): # Havia itens mas nenhuma conclusão (erro ou não detectado)
         texto_final = ("Face ao exposto e com base nos resultados obtidos, conclui-se que nos materiais examinados "
                        "não foram detectadas as substâncias Cannabis sativa L. (maconha) ou Cocaína.")
    else: # Nenhum item foi submetido a exame
         texto_final = "Não houve material submetido a exame para elaboração de conclusão."

    adicionar_paragrafo(doc, texto_final, align='justify', style='Normal')


def adicionar_custodia_material(doc, lacre):
    """Adiciona a seção '7 CUSTÓDIA DO MATERIAL'."""
    adicionar_paragrafo(doc, "7 CUSTÓDIA DO MATERIAL", style='TituloPrincipal')
    # Seção sobre Contraprova
    adicionar_paragrafo(doc, "7.1 Contraprova:", style='TituloSecundario')
    texto_contraprova = ("A(s) amostra(s) para eventual contraprova referente(s) ao(s) material(is) "
                      "examinado(s) encontra(m)-se devidamente acondicionada(s) e lacrada(s) "
                      f"neste Instituto sob o lacre nº {lacre if lacre else '_____________'}" # Espaço se não preenchido
                      ", à disposição da autoridade competente, em conformidade com a legislação vigente "
                      "e normas internas sobre cadeia de custódia.")
                      # Adicionar referência à Portaria de custódia específica de GO, se conhecida
                      # Ex: "(conforme Portaria nº XXXX/YYYY - SPTC/GO)."
    adicionar_paragrafo(doc, texto_contraprova, style='Normal', align='justify')

    # Opcional: Seção sobre Descarte do Material Remanescente
    # adicionar_paragrafo(doc, "7.2 Material Remanescente:", style='TituloSecundario')
    # texto_descarte = "O material remanescente não aproveitado para contraprova será encaminhado para descarte seguro..."
    # adicionar_paragrafo(doc, texto_descarte, style='Normal', align='justify')


def adicionar_referencias(doc, subitens_cannabis, subitens_cocaina):
    """Adiciona a seção 'REFERÊNCIAS'."""
    adicionar_paragrafo(doc, "REFERÊNCIAS", style='TituloPrincipal')
    # Referências com parágrafos separados
    adicionar_paragrafo(doc,
        "BRASIL. Ministério da Saúde. Secretaria de Vigilância Sanitária. Portaria nº 344, de 12 de maio de 1998. "
        "Aprova o Regulamento Técnico sobre substâncias e medicamentos sujeitos a controle especial. "
        "Diário Oficial da União, Brasília, DF, 19 maio 1998 e suas atualizações.",
        style='Normal', align='justify')
    adicionar_paragrafo(doc,
        "GOIÁS. Secretaria de Estado da Segurança Pública. [Inserir Portaria relevante de Goiás sobre custódia/descarte de drogas, se houver e aplicável, ex: Portaria nº 0003/2019/SSP]. "
        "Diário Oficial do Estado de Goiás, Goiânia, GO, [data da publicação].",
        style='Normal', align='justify')
    adicionar_paragrafo(doc,
        "SCIENTIFIC WORKING GROUP FOR THE ANALYSIS OF SEIZED DRUGS (SWGDRUG). Recommendations. Version 8.0. June 2019. "
        "Disponível em: <http://www.swgdrug.org>. Acesso em: [Manter data se relevante ou remover].",
        style='Normal', align='justify')

    # Referências condicionais (UNODC)
    if subitens_cannabis:
        adicionar_paragrafo(doc,
            "UNITED NATIONS OFFICE ON DRUGS AND CRIME (UNODC). Laboratory and Scientific Section. "
            "Recommended Methods for the Identification and Analysis of Cannabis and Cannabis Products. New York: UNODC, 2009.", # Verificar ano correto se necessário
            style='Normal', align='justify')
    if subitens_cocaina:
        adicionar_paragrafo(doc,
            "UNITED NATIONS OFFICE ON DRUGS AND CRIME (UNODC). Laboratory and Scientific Section. "
            "Recommended Methods for the Identification and Analysis of Cocaine in Seized Materials. New York: UNODC, 2012.",
            style='Normal', align='justify')
    # Adicionar outras referências pertinentes, se necessário


def adicionar_encerramento_assinatura(doc):
    """Adiciona a frase de encerramento, data, local e a assinatura do perito."""
    adicionar_paragrafo(doc, "\nÉ o laudo. Nada mais havendo a lavrar, encerra-se o presente.", style='Normal', align='justify')

    brasilia_tz = pytz.timezone('America/Sao_Paulo')
    hoje = datetime.now(brasilia_tz)
    # Usar o dicionário de meses
    data_formatada = f"Goiânia, {hoje.day} de {meses_portugues[hoje.month]} de {hoje.year}."

    # Adicionar parágrafo vazio para espaçamento antes da data
    doc.add_paragraph() # Adiciona espaço em branco
    adicionar_paragrafo(doc, data_formatada, align='center', style='Normal') # Centralizar local e data

    # Adicionar espaço para assinatura (vários parágrafos vazios)
    doc.add_paragraph()
    doc.add_paragraph()

    # Assinatura - Manter centralizado
    # Pode usar uma tabela de 1 célula para centralizar melhor a linha, ou só texto
    adicionar_paragrafo(doc, "________________________________________", align='center', style='Normal') # Linha de assinatura
    adicionar_paragrafo(doc, "Daniel Chendes Lima", align='center', style='Normal', bold=True) # Nome em negrito
    adicionar_paragrafo(doc, "Perito Criminal - SPTC/GO", align='center', style='Normal') # Cargo e Instituição
    adicionar_paragrafo(doc, "Matrícula nº [Sua Matrícula]", align='center', style='Normal') # Adicionar matrícula

    # Informação sobre assinatura digital (se aplicável)
    # doc.add_paragraph()
    # adicionar_paragrafo(doc, "(Laudo assinado digitalmente conforme MP nº 2.200-2/2001)", align='center', style='Normal', size=9, italic=True)


def aplicar_italico_especifico(doc):
    """Aplica estilo itálico a termos científicos específicos no documento."""
    termos_italico = [
        'Cannabis sativa',
        'Cannabis sativa L.',
        # Adicionar outros nomes científicos se necessário (Ex: Erythroxylum coca)
    ]
    expressoes_latinas = ['et al.', 'i.e.', 'e.g.', 'supra', 'infra', 'in vitro', 'in vivo']
    termos_completos = termos_italico + expressoes_latinas

    # Abordagem mais simples: iterar por runs e aplicar itálico se o texto exato for encontrado
    # Limitação: Pode não pegar termos quebrados entre runs.
    for paragraph in doc.paragraphs:
         # Para evitar modificar a lista enquanto itera, trabalhe com índices ou cópia
         # A complexidade de reescrever o parágrafo (como na versão anterior) pode ser evitada
         # se a formatação original das runs não for crucial de manter além do itálico.
         for i in range(len(paragraph.runs)):
             run = paragraph.runs[i]
             for termo in termos_completos:
                 # Verifica se o termo está EXATAMENTE no texto da run (ou começando/terminando nela)
                 # Esta abordagem é mais segura mas menos flexível que a anterior
                 if termo in run.text:
                      # Se o termo exato está na run, aplica itálico à run inteira
                      # CUIDADO: Isso colocará toda a run em itálico, mesmo que contenha outras palavras.
                      # Para precisão, a reescrita do parágrafo é melhor.
                      # Vamos tentar a reescrita simplificada:
                      if termo == run.text.strip(): # Se a run contém APENAS o termo
                           run.italic = True
                      # Caso contrário (termo no meio da run), a lógica fica complexa sem reescrever.
                      # Por simplicidade aqui, vamos aplicar à run inteira se o termo estiver contido.
                      # run.italic = True # <-- Descomente esta linha para a abordagem mais simples (menos precisa)

    # Tentar a abordagem de reescrita (mais complexa mas precisa)
    # Nota: Esta reescrita perde formatação original das runs (negrito, cor, etc)
    # que não seja explicitamente reaplicada.
    for paragraph in doc.paragraphs:
        texto_completo = paragraph.text
        if not any(termo in texto_completo for termo in termos_completos):
            continue # Pula parágrafos sem os termos

        # Salva alinhamento e estilo originais
        original_alignment = paragraph.alignment
        original_style = paragraph.style

        # Limpa o parágrafo original
        paragraph.clear()
        paragraph.text = "" # Garante que está vazio

        # Restaura alinhamento e estilo
        paragraph.alignment = original_alignment
        paragraph.style = original_style


        current_pos = 0
        while current_pos < len(texto_completo):
            found_at = -1
            term_found = None

            # Encontra a *próxima* ocorrência de qualquer termo
            for termo in termos_completos:
                pos = texto_completo.find(termo, current_pos)
                if pos != -1:
                    if term_found is None or pos < found_at:
                        found_at = pos
                        term_found = termo

            if term_found:
                # Adiciona o texto *antes* do termo (se houver)
                if found_at > current_pos:
                    run_normal = paragraph.add_run(texto_completo[current_pos:found_at])
                    # Aqui, idealmente, copiaríamos a formatação da run original correspondente
                    # run_normal.bold = ... etc.

                # Adiciona o termo encontrado em itálico
                run_italic = paragraph.add_run(term_found)
                run_italic.italic = True
                # Copiar outras formatações se necessário

                # Atualiza a posição para depois do termo encontrado
                current_pos = found_at + len(term_found)
            else:
                # Nenhum termo mais encontrado, adiciona o resto do texto
                if current_pos < len(texto_completo):
                    run_normal = paragraph.add_run(texto_completo[current_pos:])
                    # Copiar formatação original
                break # Terminou o parágrafo


def gerar_laudo_docx(dados_laudo):
    """Gera o laudo completo em formato docx."""
    document = Document()
    configurar_estilos(document) # Aplica estilos primeiro
    configurar_pagina(document) # Configura margens e tamanho
    adicionar_cabecalho_rodape(document) # Adiciona header/footer

    # Adicionar seções na ordem correta
    adicionar_preambulo(document, dados_laudo)
    subitens_cannabis, subitens_cocaina = adicionar_material_recebido(document, dados_laudo)
    adicionar_objetivo_exames(document)
    adicionar_exames(document, subitens_cannabis, subitens_cocaina)
    adicionar_resultados(document, subitens_cannabis, subitens_cocaina, dados_laudo)
    adicionar_conclusao(document, dados_laudo)
    adicionar_custodia_material(document, dados_laudo.get('lacre', ''))
    adicionar_referencias(document, subitens_cannabis, subitens_cocaina)
    adicionar_encerramento_assinatura(document)

    # Aplicar formatação final (itálico)
    aplicar_italico_especifico(document)

    return document

# --- Interface Streamlit ---
def main():
    st.set_page_config(layout="wide", page_title="Gerador de Laudo Pericial")

    # --- Cabeçalho com Logo, Título, Data e Hora ---
    col1, col2, col3 = st.columns([1, 4, 2]) # Ajustar proporções conforme necessidade

    with col1:
        # Tentar carregar a imagem localmente primeiro, depois URL
        logo_path = "logo_policia_cientifica.png" # Nome do arquivo local
        try:
            st.image(logo_path, width=150)
        except FileNotFoundError:
            # Fallback para URL se o arquivo local não for encontrado
            st.image("https://www.policiacientifica.go.gov.br/wp-content/uploads/2021/08/logomarca-branca-menor.png", width=150)
        except Exception as e:
             st.warning(f"Não foi possível carregar o logo: {e}")

    with col2:
        st.title("Gerador de Laudo Pericial")
        st.caption("Identificação de Drogas e Substâncias Correlatas - SPTC/GO") # Subtítulo

    with col3:
        # Container para Data/Hora
        data_hora_placeholder = st.empty()

        # Função para atualizar a data/hora (será chamada a cada rerun)
        def atualizar_data_hora():
            try:
                 brasilia_tz = pytz.timezone('America/Sao_Paulo')
                 now = datetime.now(brasilia_tz)
                 # Usar dicionários para garantir PT-BR
                 dia_semana = dias_semana_portugues[now.weekday()]
                 mes = meses_portugues[now.month]
                 # Formato: Quinta-feira, 17 de abril de 2025 | 12:40:10 (GMT-3)
                 data_formatada = f"{dia_semana}, {now.day} de {mes} de {now.year}"
                 hora_formatada = now.strftime("%H:%M:%S")
                 # Usar HTML para alinhar à direita e talvez estilizar
                 # Usando st.markdown para formatar melhor
                 data_hora_placeholder.markdown(
                     f"""
                     <div style="text-align: right; font-size: 0.9em; color: #AAAAAA; line-height: 1.2;">
                         <span>{data_formatada}</span><br>
                         <span>{hora_formatada} (GMT-3)</span>
                     </div>
                     """,
                     unsafe_allow_html=True
                 )
            except Exception as e:
                 # Em caso de erro (ex: pytz não instalado), mostra um fallback
                 now = datetime.now()
                 fallback_str = now.strftime("%d/%m/%Y %H:%M:%S")
                 data_hora_placeholder.markdown(
                     f"""
                     <div style="text-align: right; font-size: 0.9em; color: #FF5555; line-height: 1.2;">
                         <span>{fallback_str} (Horário Local)</span><br>
                         <span style="font-size: 0.8em;">Erro Timezone: {e}</span>
                     </div>
                     """,
                      unsafe_allow_html=True
                 )

        # Chamada inicial para exibir a hora imediatamente
        atualizar_data_hora()

    st.markdown("---") # Linha divisória

    # --- Coleta de Dados para o Laudo ---
    st.header("Informações Gerais do Laudo")

    # Usar estado da sessão para persistir dados entre reruns
    if 'dados_laudo' not in st.session_state:
        st.session_state.dados_laudo = {
            'autoridade': 'Autoridade Policial',
            'num_doc': 'Ofício nº ____/____',
            'boletim': 'BO nº ____/____',
            'lacre': '',
            'itens': [],
            'imagem': None # Armazena o objeto UploadedFile
        }

    # Inputs para informações gerais (preâmbulo)
    st.session_state.dados_laudo['autoridade'] = st.text_input(
        "Autoridade Requisitante",
        value=st.session_state.dados_laudo['autoridade']
    )
    col_doc, col_bo = st.columns(2)
    with col_doc:
        st.session_state.dados_laudo['num_doc'] = st.text_input(
            "Nº Documento Requisição (Ofício, etc.)",
            value=st.session_state.dados_laudo['num_doc']
        )
    with col_bo:
        st.session_state.dados_laudo['boletim'] = st.text_input(
            "Nº Boletim de Ocorrência (ou similar)",
            value=st.session_state.dados_laudo['boletim']
        )

    st.session_state.dados_laudo['lacre'] = st.text_input(
        "Número do Lacre da Contraprova (se houver)",
         value=st.session_state.dados_laudo['lacre']
         )

    st.markdown("---")
    st.header("Descrição dos Itens Recebidos")

    numero_itens = st.number_input(
        "Número de itens a descrever",
        min_value=0, # Permitir 0 itens
        value=max(0, len(st.session_state.dados_laudo.get('itens', []))), # Inicia com o número atual de itens
        step=1,
        key="num_itens_input", # Chave para evitar reset inesperado
        help="Quantos tipos diferentes de material/embalagem foram recebidos?"
        )

    # Ajusta a lista de itens no estado da sessão para corresponder ao número desejado
    # Garante que 'itens' seja uma lista
    if not isinstance(st.session_state.dados_laudo.get('itens'), list):
        st.session_state.dados_laudo['itens'] = []

    current_num_itens = len(st.session_state.dados_laudo['itens'])
    if numero_itens > current_num_itens:
        for _ in range(numero_itens - current_num_itens):
            # Adiciona um novo item com valores padrão
            st.session_state.dados_laudo['itens'].append({
                'qtd': 1,
                'tipo_mat': list(TIPOS_MATERIAL_BASE.keys())[0], # Default para o primeiro tipo
                'emb': list(TIPOS_EMBALAGEM_BASE.keys())[0], # Default para o primeiro tipo
                'cor_emb': None, # Default sem cor
                'ref': '', # Referência vazia por padrão
                'pessoa': '' # Pessoa vazia por padrão
            })
    elif numero_itens < current_num_itens:
        # Remove itens excedentes do final da lista
        st.session_state.dados_laudo['itens'] = st.session_state.dados_laudo['itens'][:numero_itens]

    # Exibe os campos para cada item
    if numero_itens > 0: # Só mostra a seção de itens se houver pelo menos 1
        for i in range(numero_itens):
            # Usa um container para agrupar visualmente cada item
            with st.container():
                st.subheader(f"Item {i + 1}")
                item_key_prefix = f"item_{i}_" # Prefixo único para widgets deste item

                cols_item1 = st.columns([1, 3, 3]) # Qtd | Material | Embalagem
                with cols_item1[0]:
                    # Garante que qtd existe e é número
                    if not isinstance(st.session_state.dados_laudo['itens'][i].get('qtd'), int):
                         st.session_state.dados_laudo['itens'][i]['qtd'] = 1
                    st.session_state.dados_laudo['itens'][i]['qtd'] = st.number_input(
                        f"Qtd", min_value=1,
                        value=st.session_state.dados_laudo['itens'][i]['qtd'],
                        step=1, key=item_key_prefix + "qtd"
                        )
                with cols_item1[1]:
                    # Garante que tipo_mat existe
                    if 'tipo_mat' not in st.session_state.dados_laudo['itens'][i]:
                        st.session_state.dados_laudo['itens'][i]['tipo_mat'] = list(TIPOS_MATERIAL_BASE.keys())[0]
                    st.session_state.dados_laudo['itens'][i]['tipo_mat'] = st.selectbox(
                        f"Material", options=list(TIPOS_MATERIAL_BASE.keys()),
                        format_func=lambda x: f"{x} ({TIPOS_MATERIAL_BASE.get(x, 'Desconhecido')})",
                        index=list(TIPOS_MATERIAL_BASE.keys()).index(st.session_state.dados_laudo['itens'][i]['tipo_mat']),
                        key=item_key_prefix + "tipo_mat"
                        )
                with cols_item1[2]:
                    # Garante que emb existe
                    if 'emb' not in st.session_state.dados_laudo['itens'][i]:
                        st.session_state.dados_laudo['itens'][i]['emb'] = list(TIPOS_EMBALAGEM_BASE.keys())[0]
                    st.session_state.dados_laudo['itens'][i]['emb'] = st.selectbox(
                        f"Embalagem", options=list(TIPOS_EMBALAGEM_BASE.keys()),
                        format_func=lambda x: f"{x} ({TIPOS_EMBALAGEM_BASE.get(x, 'Desconhecida')})",
                        index=list(TIPOS_EMBALAGEM_BASE.keys()).index(st.session_state.dados_laudo['itens'][i]['emb']),
                        key=item_key_prefix + "emb"
                        )

                cols_item2 = st.columns([1, 2, 2]) # Cor | Ref | Pessoa
                with cols_item2[0]:
                    emb_selecionada = st.session_state.dados_laudo['itens'][i]['emb']
                    if emb_selecionada in ['pl', 'pa', 'e', 'z']: # Embalagens que podem ter cor
                        # Garante que cor_emb existe (pode ser None)
                        if 'cor_emb' not in st.session_state.dados_laudo['itens'][i]:
                             st.session_state.dados_laudo['itens'][i]['cor_emb'] = None

                        cor_options = [''] + list(CORES_FEMININO_EMBALAGEM.keys())
                        try:
                            current_cor = st.session_state.dados_laudo['itens'][i]['cor_emb']
                            # +1 por causa do '' no início
                            default_cor_index = cor_options.index(current_cor) if current_cor in cor_options else 0
                        except ValueError:
                            default_cor_index = 0 # Se cor atual inválida, default para ''

                        selected_cor = st.selectbox(
                            f"Cor Emb.", options=cor_options,
                            format_func=lambda x: f"{x} ({CORES_FEMININO_EMBALAGEM.get(x, 'N/A')})" if x else "N/A",
                            index=default_cor_index, key=item_key_prefix + "cor_emb"
                            )
                        # Atualiza para None se selecionou a opção vazia ('')
                        st.session_state.dados_laudo['itens'][i]['cor_emb'] = selected_cor if selected_cor else None
                    else:
                        # Se a embalagem não tem cor, desabilita e mostra N/A
                        st.session_state.dados_laudo['itens'][i]['cor_emb'] = None
                        st.text_input(f"Cor Emb.", value="N/A", disabled=True, key=item_key_prefix + "cor_emb_disabled")

                with cols_item2[1]:
                     if 'ref' not in st.session_state.dados_laudo['itens'][i]: st.session_state.dados_laudo['itens'][i]['ref'] = ''
                     st.session_state.dados_laudo['itens'][i]['ref'] = st.text_input(
                        f"Ref. Constatação",
                        value=st.session_state.dados_laudo['itens'][i]['ref'],
                        key=item_key_prefix + "ref", help="Número do subitem no Laudo de Constatação, se houver."
                        )
                with cols_item2[2]:
                    if 'pessoa' not in st.session_state.dados_laudo['itens'][i]: st.session_state.dados_laudo['itens'][i]['pessoa'] = ''
                    st.session_state.dados_laudo['itens'][i]['pessoa'] = st.text_input(
                        f"Pessoa Relacionada",
                        value=st.session_state.dados_laudo['itens'][i]['pessoa'],
                        key=item_key_prefix + "pessoa", help="Nome da pessoa associada a este item (opcional)."
                        )

                # Linha divisória entre os itens (exceto após o último)
                if i < numero_itens - 1:
                    st.markdown("---")
    else:
        st.info("Adicione pelo menos 1 item para descrever.")

    st.markdown("---") # Linha divisória após a seção de itens

    # --- Adicionar Imagem (NOVA POSIÇÃO) ---
    with st.expander("Adicionar Foto do Material Recebido (Opcional)", expanded=False):
        uploaded_image = st.file_uploader(
            "Selecione um arquivo de imagem...",
            type=["png", "jpg", "jpeg", "bmp", "gif", "tiff"],
            key="image_uploader", # Chave única para o uploader
            label_visibility="collapsed" # Esconde o label padrão
            )

        # Lógica para lidar com o upload e o estado da sessão
        if uploaded_image is not None:
            # Se um NOVO arquivo foi carregado, atualiza o estado da sessão
            st.session_state.dados_laudo['imagem'] = uploaded_image
            st.image(uploaded_image, caption="Pré-visualização da nova imagem.", width=300)
        elif st.session_state.dados_laudo.get('imagem') is not None:
             # Se NÃO há novo upload MAS JÁ EXISTE imagem no estado, mostra a existente
             st.image(st.session_state.dados_laudo['imagem'], caption="Pré-visualização da imagem carregada.", width=300)
             # Botão para remover a imagem existente
             if st.button("Remover Imagem", key="remove_image_button"):
                  st.session_state.dados_laudo['imagem'] = None
                  # Limpa o estado do file_uploader explicitamente se necessário (pode não ser preciso)
                  # st.session_state.image_uploader = None
                  st.rerun() # Força o rerender da UI sem a imagem


    st.markdown("---") # Linha divisória antes da geração

    # --- Geração e Download ---
    st.header("Gerar Documento")

    if st.button("Gerar Laudo (.docx)", type="primary", use_container_width=True):
        with st.spinner("Gerando documento Word... Por favor, aguarde."):
            try:
                # Passa os dados do estado da sessão para a função geradora
                # Certifica-se de que 'imagem' contém o objeto UploadedFile se existir
                doc = gerar_laudo_docx(st.session_state.dados_laudo)

                # Salvar o documento em um buffer de memória
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0) # Volta ao início do buffer para leitura

                # Oferecer para download
                # Gerar nome de arquivo sugerido (ex: Laudo_BO_ZZZ-AAAA.docx)
                nome_arquivo_sugerido = "Laudo_Pericial"
                bo_num = st.session_state.dados_laudo.get('boletim', '').replace('/', '-').replace(' ', '_')
                # Remove caracteres inválidos para nome de arquivo (básico)
                bo_num = re.sub(r'[\\/*?:"<>|]', "", bo_num)
                if bo_num and bo_num != 'BO_nº____-____':
                     nome_arquivo_sugerido += f"_{bo_num}"
                nome_arquivo_sugerido += f"_{datetime.now().strftime('%Y%m%d')}.docx" # Adiciona data


                st.success("Laudo gerado com sucesso!")
                st.download_button(
                    label="✔️ Baixar Laudo (DOCX)",
                    data=buffer,
                    file_name=nome_arquivo_sugerido,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                 st.error(f"❌ Erro ao gerar o laudo DOCX:")
                 st.exception(e) # Mostra o traceback formatado do Streamlit


    # Adicionar um pequeno rodapé na interface
    st.markdown("---")
    st.caption("Gerador de Laudo Pericial v2.5 :: Superintendência de Polícia Técnico-Científica :: Goiás")


if __name__ == "__main__":
    # Todas as definições de funções e constantes devem estar antes desta linha
    main()
